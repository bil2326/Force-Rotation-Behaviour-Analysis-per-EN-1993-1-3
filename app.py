import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from io import StringIO, BytesIO
from scipy.signal import savgol_filter
import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

st.set_page_config(page_title="Analyse Courbe F–δ", layout="wide", page_icon="📊")

st.markdown("""
    <h2 style='color:#1e3a5f;'>📊 Analyse Courbe F–δ — Figure A.9 (EN 1993-1-3)</h2>
    <p style='color:#475569;'>Détermination automatique de δ_el et δ_pl par interpolation linéaire</p>
    <hr>
""", unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.header("⚙️ Paramètres d'essai")
    s = st.number_input("Portée s (mm)", value=3600.0, step=1.0)
    e = st.number_input("Distance appui e (mm)", value=800.0, step=1.0)

    st.markdown("---")
    st.markdown("""
**Formule (Éq. A.4a) :**

`θ = 2(δ_pl − δ_e − δ_el) / (0.5·s − e)`

avec :

`δ_e = δ_e_montante`

- **δ_el** → branche montante F–δ
- **δ_pl** → branche descendante F–δ
- **δ_e** → branche montante F–δ_e uniquement

---
**Méthode K_φ :**

K_φ = M_last / θ_last

_M_last = dernier point où θ > 0 (limite domaine valide Éq. A.4a)_
""")

# ════════════════════════════════════════════════════════════════════════════
# FONCTIONS
# ════════════════════════════════════════════════════════════════════════════

def parse_paste(text):
    """Parse un collage Excel 2 colonnes → DataFrame [F, val]."""
    try:
        df = pd.read_csv(StringIO(text.strip()), sep=r"\t|;", engine="python",
                         header=None, on_bad_lines="skip")
        df = df.iloc[:, :2].copy()
        df.columns = ["F", "val"]
        df["F"]   = pd.to_numeric(df["F"].astype(str).str.replace(",", "."), errors="coerce")
        df["val"] = pd.to_numeric(df["val"].astype(str).str.replace(",", "."), errors="coerce")
        df = df.dropna().reset_index(drop=True)
        return df if len(df) >= 2 else None
    except:
        return None


def interp_branch(df, F_col, val_col, F_target):
    """
    Interpolation linéaire robuste sur UNE branche.
    - Regroupe par F (moyenne) pour éliminer les doublons
    - Trie par F croissant (requis par np.interp)
    - Retourne None si F_target hors plage
    """
    tmp = df[[F_col, val_col]].copy()
    tmp.columns = ["F", "v"]
    tmp = tmp.groupby("F", as_index=False)["v"].mean()
    tmp = tmp.sort_values("F").reset_index(drop=True)

    f = tmp["F"].values
    v = tmp["v"].values

    if f.min() <= F_target <= f.max():
        return float(np.interp(F_target, f, v))
    return None


def split_branches(df, f_col, val_col):
    """
    Sépare un DataFrame en branche montante et descendante.
    Retourne (fmax_idx, ascending, descending).
    """
    fmax_idx   = int(df[f_col].idxmax())
    ascending  = df.iloc[:fmax_idx + 1].reset_index(drop=True)
    descending = df.iloc[fmax_idx:].reset_index(drop=True)
    return fmax_idx, ascending, descending


# ════════════════════════════════════════════════════════════════════════════
# LAYOUT
# ════════════════════════════════════════════════════════════════════════════
col1, col2 = st.columns([1, 1.6])

# ════════════════════════════════════════════════════════════════════════════
# COLONNE GAUCHE — saisies
# ════════════════════════════════════════════════════════════════════════════
with col1:

    # ── 1. Courbe F–δ ────────────────────────────────────────────────────────
    st.subheader("📋 Courbe F–δ (flèche milieu)")
    st.caption("Collez vos 2 colonnes **(F | δ)** depuis Excel :")

    paste_fdelta = st.text_area(
        "fdelta", height=200,
        placeholder="F (N)\tδ (mm)\n0\t0\n500\t0.55\n1000\t1.18\n...",
        label_visibility="collapsed", key="paste_fdelta"
    )

    df_points  = None
    fmax_idx   = None
    fmax_val   = None
    ascending  = None
    descending = None

    if paste_fdelta.strip():
        df_raw = parse_paste(paste_fdelta)
        if df_raw is not None:
            df_raw.columns = ["F", "delta"]
            df_points = df_raw

            fmax_idx, ascending, descending = split_branches(df_points, "F", "delta")
            fmax_val = float(df_points["F"].iloc[fmax_idx])

            def blabel(i):
                if i < fmax_idx:  return "↗ montante"
                if i == fmax_idx: return "⭐ F_max"
                return "↘ descendante"

            st.success(
                f"✅ **{len(df_points)} points** — "
                f"F_max = **{fmax_val:.3f} kN** | "
                f"montante : {fmax_idx + 1} pts | "
                f"descendante : {len(df_points) - fmax_idx} pts"
            )
            with st.expander("👁 Aperçu F–δ"):
                dv = df_points.copy()
                dv["F (kN)"]  = dv["F"].round(4)
                dv["Branche"] = [blabel(i) for i in range(len(df_points))]
                dv = dv.drop(columns=["F"])
                dv.index += 1
                st.dataframe(
                    dv.rename(columns={"delta": "δ (mm)"}),
                    use_container_width=True, height=200
                )
        else:
            st.error("❌ Données invalides — vérifiez vos colonnes.")

    st.markdown("---")

    # ── 2. Courbe F–δ_e ──────────────────────────────────────────────────────
    st.subheader("📋 Courbe F–δ_e (flèche appuis)")

    mode_de = st.radio(
        "Mode δ_e :",
        ["Valeur constante", "Courbe mesurée (Excel)"],
        horizontal=True
    )

    df_de         = None
    de_cst        = 0.0
    ascending_de  = None
    descending_de = None
    fmax_idx_de   = None

    if mode_de == "Valeur constante":
        de_cst = st.number_input(
            "δ_e (mm) — valeur unique pour tous les F",
            value=0.0, step=0.001, format="%.4f"
        )
        st.info("δ_e = constante appliquée à tous les niveaux de charge.")

    else:
        st.caption("Collez vos 2 colonnes **(F | δ_e)** depuis Excel :")
        paste_de = st.text_area(
            "de_paste", height=180,
            placeholder="F (N)\tδ_e (mm)\n0\t0\n500\t0.01\n1000\t0.02\n...",
            label_visibility="collapsed", key="paste_de"
        )

        if paste_de.strip():
            df_de_raw = parse_paste(paste_de)
            if df_de_raw is not None:
                df_de_raw.columns = ["F", "de"]
                df_de = df_de_raw

                fmax_idx_de, ascending_de, descending_de = split_branches(df_de, "F", "de")

                st.success(
                    f"✅ **{len(df_de)} points** δ_e importés — "
                    f"F : {df_de['F'].min():.3f} → {df_de['F'].max():.3f} kN | "
                    f"montante : {fmax_idx_de + 1} pts | "
                    f"descendante : {len(df_de) - fmax_idx_de} pts"
                )
                with st.expander("👁 Aperçu F–δ_e"):
                    dv2 = df_de.copy()
                    dv2["F (kN)"] = dv2["F"].round(4)
                    dv2 = dv2.drop(columns=["F"])
                    dv2.index += 1
                    st.dataframe(
                        dv2.rename(columns={"de": "δ_e (mm)"}),
                        use_container_width=True, height=180
                    )
            else:
                st.error("❌ Données δ_e invalides.")

    st.markdown("---")

    # ── 3. Niveaux F ─────────────────────────────────────────────────────────
    st.subheader("🎯 Niveaux F à analyser")
    st.caption(f"Entrez les valeurs en **kN** (un F par ligne) :")

    f_input = st.text_area(
        "fvals", height=140,
        placeholder="10\n20\n30\n35\n39",
        label_visibility="collapsed", key="f_input_area"
    )

    bc1, bc2 = st.columns(2)
    with bc1:
        if st.button("⚡ Auto 10 niveaux", use_container_width=True):
            if df_points is not None:
                st.session_state["f_auto"] = "\n".join(
                    str(round(fmax_val * (i + 1) / 11, 3)) for i in range(10)
                )
    with bc2:
        if st.button("📐 10%…95% F_max", use_container_width=True):
            if df_points is not None:
                st.session_state["f_auto"] = "\n".join(
                    str(round(fmax_val * p / 100, 3))
                    for p in [10, 20, 30, 40, 50, 60, 70, 80, 90, 95]
                )

    if "f_auto" in st.session_state and not f_input.strip():
        f_input = st.session_state["f_auto"]
        st.text_area(
            "(auto)", value=f_input, height=120,
            label_visibility="collapsed", disabled=True
        )

    run = st.button("▶ Analyser", type="primary", use_container_width=True)

# ════════════════════════════════════════════════════════════════════════════
# ANALYSE
# ════════════════════════════════════════════════════════════════════════════
results_df = None

if run:
    if df_points is None:
        st.error("❌ Importez d'abord la courbe F–δ.")
    elif not f_input.strip():
        st.error("❌ Entrez au moins un niveau de charge F.")
    else:
        raw_f  = f_input.replace("\t", "\n").replace(";", "\n").split("\n")
        f_list = []
        for v in raw_f:
            try:
                fv = float(v.strip().replace(",", "."))
                if 0 < fv < fmax_val:
                    f_list.append(fv)
            except:
                pass

        if not f_list:
            st.error(f"❌ Aucune valeur F valide (0 < F < {fmax_val:.3f} kN).")
        else:
            rows, errs = [], []

            for F in sorted(set(f_list)):

                # δ_el — branche MONTANTE de F–δ
                del_val = interp_branch(ascending,  "F", "delta", F)
                # δ_pl — branche DESCENDANTE de F–δ
                dpl_val = interp_branch(descending, "F", "delta", F)

                if del_val is None or dpl_val is None:
                    errs.append(round(F, 4))
                    continue

                if mode_de == "Valeur constante":
                    de_val_el = de_cst
                    de_val_pl = de_cst
                    de_moy    = de_cst
                else:
                    if ascending_de is not None and descending_de is not None:
                        de_val_el = interp_branch(ascending_de,  "F", "de", F)
                        de_val_pl = interp_branch(descending_de, "F", "de", F)
                        if de_val_el is None:
                            st.warning(f"⚠️ F={F:.3f} kN hors plage δ_e montante → 0 utilisé.")
                            de_val_el = 0.0
                        if de_val_pl is None:
                            st.warning(f"⚠️ F={F:.3f} kN hors plage δ_e descendante → 0 utilisé.")
                            de_val_pl = 0.0
                        de_moy = de_val_el
                    else:
                        de_val_el = 0.0
                        de_val_pl = 0.0
                        de_moy    = 0.0

                if (dpl_val - del_val) < de_moy:
                    errs.append(round(F, 4))
                    continue

                denom = 0.5 * s - e
                if abs(denom) < 1e-12:
                    st.error("❌ Dénominateur nul : vérifiez s et e.")
                    break

                theta = (2.0 * (dpl_val - de_moy - del_val)) / denom
                M_kNm = F * s / 4.0 / 1000.0

                rows.append({
                    "F (kN)":                  round(F, 4),
                    "F / F_max (%)":           round(F / fmax_val * 100, 1),
                    "δ_el (mm)":               round(del_val, 4),
                    "δ_pl (mm)":               round(dpl_val, 4),
                    "δ_e montante (mm)":       round(de_val_el, 4),
                    "δ_e descendante (mm)":    round(de_val_pl, 4),
                    "δ_e moy (mm)":            round(de_moy, 4),
                    "δ_pl − δ_el (mm)":        round(dpl_val - del_val, 4),
                    "θ (rad) [Éq. A.4a]":      round(theta, 6),
                    "M = F·s/4 (kN·m)":        round(M_kNm, 4),
                })

            if rows:
                results_df = pd.DataFrame(rows)
                st.session_state["results_df"]  = results_df
                st.session_state["results_pts"] = [
                    {"F": r["F (kN)"], "del": r["δ_el (mm)"], "dpl": r["δ_pl (mm)"]}
                    for r in rows
                ]
            if errs:
                st.warning(
                    f"⚠️ Points ignorés (kN) : {errs}  \n"
                    f"_(hors plage d'interpolation **ou** δ_pl − δ_el < δ_e_montante)_"
                )

elif "results_df" in st.session_state:
    results_df = st.session_state["results_df"]

# ════════════════════════════════════════════════════════════════════════════
# COLONNE DROITE — graphiques + résultats
# ════════════════════════════════════════════════════════════════════════════
with col2:

    # ── Graphique F–δ ─────────────────────────────────────────────────────
    st.subheader("📈 Courbe F–δ")
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.set_facecolor("#f8fafc")
    fig.patch.set_facecolor("white")

    if df_points is not None:
        up   = df_points.iloc[:fmax_idx + 1]
        down = df_points.iloc[fmax_idx:]

        ax.plot(up["delta"],   up["F"],   color="#2563eb", lw=2.2, label="Branche montante")
        ax.plot(down["delta"], down["F"], color="#ea580c", lw=2.2, label="Branche descendante")
        ax.scatter(
            df_points["delta"].iloc[fmax_idx],
            fmax_val,
            color="#dc2626", zorder=5, s=70,
            label=f"F_max = {fmax_val:.3f} kN"
        )

        for r in st.session_state.get("results_pts", []):
            ax.axhline(r["F"], color="#cbd5e1", lw=0.7, ls="--", zorder=1)
            ax.scatter(r["del"], r["F"], color="#2563eb", marker="^", s=80, zorder=6)
            ax.scatter(r["dpl"], r["F"], color="#ea580c", marker="^", s=80, zorder=6)

        p1 = mpatches.Patch(color="#2563eb", label="▲ δ_el")
        p2 = mpatches.Patch(color="#ea580c", label="▲ δ_pl")
        handles, _ = ax.get_legend_handles_labels()
        ax.legend(handles=handles[:3] + [p1, p2], fontsize=8)

    ax.set_xlabel("Flèche nette δ (mm)", fontsize=10)
    ax.set_ylabel("Charge F (kN)", fontsize=10)
    ax.grid(True, alpha=0.4, lw=0.6)
    ax.spines[["top", "right"]].set_visible(False)
    st.pyplot(fig, use_container_width=True)

    # ── Graphique F–δ_e ──────────────────────────────────────────────────
    if df_de is not None and fmax_idx_de is not None:
        with st.expander("📈 Courbe F–δ_e (flèche aux appuis)"):
            fig2, ax2 = plt.subplots(figsize=(7, 3))
            up_de   = df_de.iloc[:fmax_idx_de + 1]
            down_de = df_de.iloc[fmax_idx_de:]
            ax2.plot(up_de["de"],   up_de["F"],   color="#2563eb", lw=2, label="Montante")
            ax2.plot(down_de["de"], down_de["F"], color="#ea580c", lw=2, label="Descendante")
            ax2.set_xlabel("δ_e (mm)", fontsize=10)
            ax2.set_ylabel("F (kN)", fontsize=10)
            ax2.legend(fontsize=8)
            ax2.grid(True, alpha=0.4)
            ax2.spines[["top", "right"]].set_visible(False)
            ax2.set_facecolor("#f8fafc")
            fig2.patch.set_facecolor("white")
            st.pyplot(fig2, use_container_width=True)

    # ── Résultats ─────────────────────────────────────────────────────────
    if results_df is not None:
        st.subheader("📋 Résultats")

        fmt = {
            "F (kN)":                  "{:.4f}",
            "F / F_max (%)":           "{:.1f}",
            "δ_el (mm)":               "{:.4f}",
            "δ_pl (mm)":               "{:.4f}",
            "δ_e montante (mm)":       "{:.4f}",
            "δ_e descendante (mm)":    "{:.4f}",
            "δ_e moy (mm)":            "{:.4f}",
            "δ_pl − δ_el (mm)":        "{:.4f}",
            "θ (rad) [Éq. A.4a]":      "{:.6f}",
            "M = F·s/4 (kN·m)":        "{:.4f}",
        }

        styled = (results_df.style
            .applymap(lambda _: "color:#2563eb; font-weight:bold", subset=["δ_el (mm)"])
            .applymap(lambda _: "color:#ea580c; font-weight:bold", subset=["δ_pl (mm)"])
            .applymap(lambda _: "color:#7c3aed; font-weight:bold",
                      subset=["δ_e montante (mm)", "δ_e descendante (mm)", "δ_e moy (mm)"])
            .applymap(lambda _: "color:#0f766e; font-weight:bold", subset=["M = F·s/4 (kN·m)"])
            .format(fmt)
        )
        st.dataframe(styled, use_container_width=True, hide_index=True)

        # ── Graphique M–θ ─────────────────────────────────────────────────
        st.subheader("📈 Courbe M–θ (Fig. A.10 — EN 1993-1-3)")

        st.caption(
            "Chaque clic sur **▶ Analyser** ajoute un essai. "
            "Cliquez sur **🗑 Effacer tous les essais** pour repartir de zéro."
        )

        # ── Initialisation stockage courbes ──────────────────────────────
        if "mt_curves" not in st.session_state:
            st.session_state["mt_curves"] = []

        # ── Ajout courbe courante ─────────────────────────────────────────
        if run and results_df is not None:
            curve_data = results_df[["θ (rad) [Éq. A.4a]", "M = F·s/4 (kN·m)"]].copy()
            curve_data.columns = ["theta", "M"]
            # filtre points valides
            curve_data = curve_data[curve_data["theta"] > 0].reset_index(drop=True)
            already = any(
                len(c) == len(curve_data) and
                np.allclose(c["theta"].values, curve_data["theta"].values, atol=1e-8)
                for c in st.session_state["mt_curves"]
            )
            if not already and len(curve_data) >= 2:
                st.session_state["mt_curves"].append(curve_data)

        curves = st.session_state["mt_curves"]

        COLORS = ["#2563eb","#ea580c","#16a34a","#7c3aed","#0891b2",
                  "#be185d","#b45309","#374151"]

        col_mt1, col_mt2 = st.columns([3, 1])

        with col_mt2:
            st.metric("Essais chargés", len(curves))
            if st.button("🗑 Effacer tous les essais", use_container_width=True):
                st.session_state["mt_curves"] = []
                st.session_state.pop("mt_grid", None)
                st.rerun()

        with col_mt1:
            fig3, ax3 = plt.subplots(figsize=(8, 5))
            ax3.set_facecolor("#f8fafc")
            fig3.patch.set_facecolor("white")

            if curves:
                # ── BUG CORRIGÉ : boucle for manquante ───────────────────
                for i, c in enumerate(curves):
                    col_i = COLORS[i % len(COLORS)]

                    t_vals = c["theta"].values.astype(float)
                    m_vals = c["M"].values.astype(float)

                    # suppression bruit initial
                    mask = t_vals > 1e-5
                    t_vals = t_vals[mask]
                    m_vals = m_vals[mask]

                    if len(m_vals) < 2:
                        continue

                    # lissage léger si assez de points
                    if len(m_vals) > 21:
                        m_lisse = savgol_filter(m_vals, min(21, len(m_vals) if len(m_vals) % 2 != 0 else len(m_vals) - 1), 3)
                    else:
                        m_lisse = m_vals

                    ax3.plot(t_vals, m_lisse,
                             color=col_i, lw=2.0, alpha=0.9,
                             label=f"Essai {i+1} (lissé)")

                # ── Moyenne M sur θ commun ────────────────────────────────
                theta_min = max(c["theta"].min() for c in curves)
                theta_max = min(c["theta"].max() for c in curves)

                if theta_max > theta_min:
                    theta_grid = np.linspace(theta_min, theta_max, 300)
                    M_interp_all = []

                    for c in curves:
                        t = c["theta"].values.astype(float)
                        m = c["M"].values.astype(float)

                        # tri par theta croissant
                        idx_sort = np.argsort(t)
                        t = t[idx_sort]
                        m = m[idx_sort]

                        # suppression bruit
                        mask2 = t > 1e-4
                        t = t[mask2]
                        m = m[mask2]

                        if len(t) < 2:
                            continue

                        # lissage rolling
                        m_smooth = (
                            pd.Series(m)
                            .rolling(window=min(9, len(m)), center=True)
                            .mean()
                            .bfill()
                            .ffill()
                            .to_numpy()
                        )

                        M_interp_all.append(np.interp(theta_grid, t, m_smooth))

                    if M_interp_all:
                        M_mean = np.array(M_interp_all).mean(axis=0)

                        # lissage final
                        win = min(11, len(M_mean) if len(M_mean) % 2 != 0 else len(M_mean) - 1)
                        M_smooth_final = (
                            pd.Series(M_mean)
                            .rolling(window=win, center=True)
                            .mean()
                            .bfill()
                            .ffill()
                            .to_numpy()
                        )

                        # ── Calcul K_φ : méthode M_last / θ_last ──────────
                        # K_φ basé sur M_max de la courbe lissée
                        # On prend le point (θ, M) où M est maximal.

                        M_09 = 0.9 * M_smooth_final

                        idx_max        = int(np.argmax(M_smooth_final))
                        theta_last_val = float(theta_grid[idx_max])
                        M_last_lisse   = float(M_smooth_final[idx_max])
                        M_last_09      = float(M_09[idx_max])

                        if theta_last_val > 1e-8:
                            K_phi_brut = M_last_lisse / theta_last_val
                            K_phi_exp  = M_last_09    / theta_last_val
                            theta_kphi = theta_last_val
                            M_kphi     = M_last_lisse
                            R2         = float('nan')
                        else:
                            K_phi_brut = float('nan')
                            K_phi_exp  = float('nan')
                            theta_kphi = float('nan')
                            M_kphi     = float('nan')
                            R2         = float('nan')

                        # ── tracés ────────────────────────────────────────
                        ax3.plot(theta_grid, M_smooth_final,
                                 color="#111827", lw=2.5,
                                 label="Courbe M–θ lissée", zorder=5)

                        ax3.plot(theta_grid, M_09,
                                 color="#dc2626", lw=2.0, ls="--",
                                 label="0,9 × M–θ lissée", zorder=6)

                        # stockage export
                        st.session_state["mt_grid"] = {
                            "theta":        theta_grid,
                            "M_mean_brut":  M_mean,
                            "M_mean_lisse": M_smooth_final,
                            "M_09":         M_09,
                            "K_phi_exp":    K_phi_exp,
                            "K_phi_brut":   K_phi_brut,
                            "theta_kphi":   theta_kphi,
                            "M_kphi":       M_kphi,
                        }

                        # ── Tableau résultats K_φ ─────────────────────────
                        if not np.isnan(K_phi_exp):
                            df_kphi = pd.DataFrame({
                                "Grandeur": [
                                    "M_last",
                                    "\u03B8_last",
                                    "M_last / \u03B8_last",
                                    "Facteur \u00A7A.5.2.3(6)",
                                    "K\u03C6,exp retenu",
                                ],
                                "Valeur": [
                                    f"{M_kphi:.3f}",
                                    f"{theta_kphi:.6f}",
                                    f"{K_phi_brut:.1f}",
                                    "0,9",
                                    f"{K_phi_exp:.1f}",
                                ],
                                "Unité": [
                                    "kN\u00B7m",
                                    "rad",
                                    "kN\u00B7m/rad",
                                    "\u2014",
                                    "kN\u00B7m/rad",
                                ],
                            })
                            st.dataframe(
                                df_kphi.style
                                    .apply(lambda x: [
                                        "background-color:#fef3c7; font-weight:bold"
                                        if i == len(df_kphi) - 1 else ""
                                        for i in range(len(df_kphi))
                                    ], axis=0),
                                use_container_width=True,
                                hide_index=True,
                            )

            ax3.set_xlabel("θ (rad)", fontsize=11)
            ax3.set_ylabel("M = F·s/4  (kN·m)", fontsize=11)
            ax3.set_title("Diagramme M–θ — Fig. A.10 (EN 1993-1-3)", fontsize=11)
            ax3.legend(fontsize=8, loc="upper left")
            ax3.grid(True, alpha=0.35, lw=0.6)
            ax3.spines[["top", "right"]].set_visible(False)
            st.pyplot(fig3, use_container_width=True)

        # ── Export M–θ ────────────────────────────────────────────────────
        if "mt_grid" in st.session_state and curves:
            g = st.session_state["mt_grid"]
            df_mt = pd.DataFrame({
                "θ (rad)":                  g["theta"],
                "M_mean brut (kN·m)":       g["M_mean_brut"],
                "M_lissée (kN·m)":          g["M_mean_lisse"],
                "0,9 × M_lissée (kN·m)":   g["M_09"],
            })
            buf_mt = BytesIO()
            with pd.ExcelWriter(buf_mt, engine="openpyxl") as w:
                df_mt.to_excel(w, index=False, sheet_name="M-theta grille")
                if "results_df" in st.session_state:
                    st.session_state["results_df"].to_excel(
                        w, index=False, sheet_name="Détail F-δ-θ"
                    )
            st.download_button(
                "⬇ Export M–θ (Excel)",
                buf_mt.getvalue(),
                "courbe_M_theta.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )

        st.markdown("---")

        with st.expander("🔍 Détail calcul θ (vérification)"):
            for _, row in results_df.iterrows():
                st.markdown(
                    f"**F = {row['F (kN)']:.4f} kN** → "
                    f"θ = 2×({row['δ_pl (mm)']:.4f} − {row['δ_e moy (mm)']:.4f} − {row['δ_el (mm)']:.4f}) "
                    f"/ (0.5×{s:.0f} − {e:.0f}) = **{row['θ (rad) [Éq. A.4a]']:.6f} rad**  "
                    f"_(δ_e = δ_e montante = {row['δ_e montante (mm)']:.4f} mm)_"
                )

        # ── Export résultats ──────────────────────────────────────────────
        ec1, ec2 = st.columns(2)
        with ec1:
            csv_b = results_df.to_csv(
                index=False, sep=";", decimal=","
            ).encode("utf-8-sig")
            st.download_button(
                "⬇ CSV (Excel FR)", csv_b,
                "resultats.csv", "text/csv",
                use_container_width=True
            )
        with ec2:
            buf = BytesIO()
            with pd.ExcelWriter(buf, engine="openpyxl") as w:
                results_df.to_excel(w, index=False, sheet_name="Résultats")
            st.download_button(
                "⬇ Excel (.xlsx)", buf.getvalue(),
                "resultats.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )

        st.markdown("---")

        # ════════════════════════════════════════════════════════════════════
        # COMPTE RENDU WORD
        # ════════════════════════════════════════════════════════════════════
        st.subheader("📄 Compte rendu Word")

        if st.button("📝 Générer le compte rendu (.docx)", type="primary", use_container_width=True):

            if not DOCX_OK:
                st.error("❌ Le package **python-docx** n'est pas installé.\n\n"
                         "Exécutez dans votre environnement :\n\n"
                         "```\npip install python-docx\n```")
                st.stop()

            def set_cell_bg(cell, hex_color):
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  hex_color)
                tcPr.append(shd)

            def add_heading(doc, text, level=1):
                p = doc.add_heading(text, level=level)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f) if level == 1 else RGBColor(0x2E, 0x40, 0x57)
                return p

            def add_para(doc, text, bold=False, italic=False, size=11):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(text)
                run.bold   = bold
                run.italic = italic
                run.font.size = Pt(size)
                run.font.name = "Arial"
                return p

            def fig_to_bytes(fig):
                buf = BytesIO()
                fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
                buf.seek(0)
                return buf

            doc = Document()

            # ── Marges ──
            for section in doc.sections:
                section.top_margin    = Cm(2.0)
                section.bottom_margin = Cm(2.0)
                section.left_margin   = Cm(2.5)
                section.right_margin  = Cm(2.5)

            # ── Style de base ──
            style = doc.styles["Normal"]
            style.font.name = "Arial"
            style.font.size = Pt(11)

            # ══════════════════════════════
            # PAGE DE TITRE
            # ══════════════════════════════
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("COMPTE RENDU D'ESSAI")
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
            run.font.name = "Arial"

            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run("Identification de la rigidité de connexion Kφ")
            run2.font.size = Pt(14)
            run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            run2.font.name = "Arial"

            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run3 = p3.add_run(f"Généré le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}")
            run3.font.size = Pt(10)
            run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run3.font.name = "Arial"

            doc.add_paragraph()

            # ══════════════════════════════
            # 1. PARAMÈTRES D'ESSAI
            # ══════════════════════════════
            add_heading(doc, "1. Paramètres d'essai", level=1)

            rows_data = [
                ["Portée s",               f"{s:.0f}",  "mm"],
                ["Distance appui e",        f"{e:.0f}",  "mm"],
                ["0,5·s − e",               f"{0.5*s-e:.0f}", "mm"],
                ["Formule rotation",        "Éq. A.4a — EN 1993-1-3", ""],
                ["Méthode Kφ",              "0,9 × M_max_lissé / θ(M_max)", ""],
            ]
            tbl = doc.add_table(rows=len(rows_data)+1, cols=3)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers = ["Paramètre", "Valeur", "Unité"]
            for i, h in enumerate(headers):
                cell = tbl.rows[0].cells[i]
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(h)
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_bg(cell, "2E75B6")
            for ri, row_d in enumerate(rows_data):
                for ci, val in enumerate(row_d):
                    cell = tbl.rows[ri+1].cells[ci]
                    cell.paragraphs[0].clear()
                    run = cell.paragraphs[0].add_run(val)
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    if ri % 2 == 1:
                        set_cell_bg(cell, "EBF3FB")

            doc.add_paragraph()

            # ══════════════════════════════
            # 2. COURBE F–δ
            # ══════════════════════════════
            add_heading(doc, "2. Courbe F–δ (flèche milieu)", level=1)
            add_para(doc, "La courbe ci-dessous présente la relation force–flèche mesurée au milieu de la panne, avec identification des branches montante et descendante et des niveaux de charge analysés.")

            if df_points is not None:
                fig_fd, ax_fd = plt.subplots(figsize=(8, 4.5))
                ax_fd.set_facecolor("#f8fafc")
                up_r   = df_points.iloc[:fmax_idx + 1]
                down_r = df_points.iloc[fmax_idx:]
                ax_fd.plot(up_r["delta"],   up_r["F"],   color="#2563eb", lw=2.2, label="Branche montante")
                ax_fd.plot(down_r["delta"], down_r["F"], color="#ea580c", lw=2.2, label="Branche descendante")
                ax_fd.scatter(df_points["delta"].iloc[fmax_idx], fmax_val,
                              color="#dc2626", zorder=5, s=70, label=f"F_max = {fmax_val:.3f} kN")
                for r in st.session_state.get("results_pts", []):
                    ax_fd.axhline(r["F"], color="#cbd5e1", lw=0.7, ls="--", zorder=1)
                    ax_fd.scatter(r["del"], r["F"], color="#2563eb", marker="^", s=80, zorder=6)
                    ax_fd.scatter(r["dpl"], r["F"], color="#ea580c", marker="^", s=80, zorder=6)
                ax_fd.set_xlabel("Flèche nette δ (mm)", fontsize=11)
                ax_fd.set_ylabel("Charge F (kN)", fontsize=11)
                ax_fd.set_title("Courbe F–δ — Fig. A.9 (EN 1993-1-3)", fontsize=11)
                ax_fd.legend(fontsize=9)
                ax_fd.grid(True, alpha=0.4, lw=0.6)
                ax_fd.spines[["top","right"]].set_visible(False)
                fig_fd.tight_layout()
                doc.add_picture(fig_to_bytes(fig_fd), width=Cm(14))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                plt.close(fig_fd)
                cap = doc.add_paragraph("Figure 1 — Courbe F–δ avec branches montante et descendante")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)

            doc.add_paragraph()

            # ══════════════════════════════
            # 3. COURBE F–δ_e
            # ══════════════════════════════
            if df_de is not None and fmax_idx_de is not None:
                add_heading(doc, "3. Courbe F–δe (flèche aux appuis)", level=1)
                add_para(doc, "La courbe ci-dessous présente la relation force–flèche aux appuis, utilisée pour le calcul de δe dans la formule A.4a.")

                fig_de, ax_de = plt.subplots(figsize=(8, 4))
                ax_de.set_facecolor("#f8fafc")
                up_de   = df_de.iloc[:fmax_idx_de + 1]
                down_de = df_de.iloc[fmax_idx_de:]
                ax_de.plot(up_de["de"],   up_de["F"],   color="#2563eb", lw=2, label="Montante")
                ax_de.plot(down_de["de"], down_de["F"], color="#ea580c", lw=2, label="Descendante")
                ax_de.set_xlabel("δe (mm)", fontsize=11)
                ax_de.set_ylabel("F (kN)", fontsize=11)
                ax_de.set_title("Courbe F–δe (flèche aux appuis)", fontsize=11)
                ax_de.legend(fontsize=9)
                ax_de.grid(True, alpha=0.4)
                ax_de.spines[["top","right"]].set_visible(False)
                fig_de.tight_layout()
                doc.add_picture(fig_to_bytes(fig_de), width=Cm(14))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                plt.close(fig_de)
                cap = doc.add_paragraph("Figure 2 — Courbe F–δe aux appuis")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
                doc.add_paragraph()

            # ══════════════════════════════
            # 4. COURBE M–θ
            # ══════════════════════════════
            fig_num = 3 if (df_de is not None and fmax_idx_de is not None) else 2
            add_heading(doc, f"{fig_num - 1}. Courbe M–θ et identification de Kφ", level=1)
            add_para(doc, "Le diagramme moment–rotation est obtenu par application de l'équation A.4a de l'EN 1993-1-3. La courbe lissée et la courbe réduite à 0,9 sont tracées. Kφ est calculé au point M_max de la courbe lissée.")

            if "mt_grid" in st.session_state:
                g = st.session_state["mt_grid"]
                fig_mt, ax_mt = plt.subplots(figsize=(8, 5))
                ax_mt.set_facecolor("#f8fafc")
                # courbes individuelles
                for i, c in enumerate(st.session_state.get("mt_curves", [])):
                    COLORS = ["#2563eb","#ea580c","#16a34a","#7c3aed"]
                    t_v = c["theta"].values.astype(float)
                    m_v = c["M"].values.astype(float)
                    mask = t_v > 1e-5
                    ax_mt.plot(t_v[mask], m_v[mask], color=COLORS[i % len(COLORS)],
                               lw=1.2, alpha=0.4, label=f"Essai {i+1}")
                ax_mt.plot(g["theta"], g["M_mean_lisse"],
                           color="#111827", lw=2.5, label="M–θ lissée")
                ax_mt.plot(g["theta"], g["M_09"],
                           color="#dc2626", lw=2.0, ls="--", label="0,9 × M–θ lissée")
                ax_mt.set_xlabel("θ (rad)", fontsize=11)
                ax_mt.set_ylabel("M = F·s/4  (kN·m)", fontsize=11)
                ax_mt.set_title("Diagramme M–θ — Fig. A.10 (EN 1993-1-3)", fontsize=11)
                ax_mt.legend(fontsize=9)
                ax_mt.grid(True, alpha=0.35, lw=0.6)
                ax_mt.spines[["top","right"]].set_visible(False)
                fig_mt.tight_layout()
                doc.add_picture(fig_to_bytes(fig_mt), width=Cm(14))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                plt.close(fig_mt)
                cap = doc.add_paragraph(f"Figure {fig_num} — Courbe M–θ lissée et 0,9 × M–θ")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)

            doc.add_paragraph()

            # ══════════════════════════════
            # 5. TABLEAU RÉSULTATS Kφ
            # ══════════════════════════════
            add_heading(doc, f"{fig_num}. Résultats — Identification de Kφ", level=1)
            add_para(doc, "Le tableau suivant présente les valeurs utilisées pour l'identification de la rigidité de connexion Kφ selon l'EN 1993-1-3 §A.5.2.3(6).")
            doc.add_paragraph()

            if "mt_grid" in st.session_state:
                g = st.session_state["mt_grid"]
                kphi_rows = [
                    ["M_max lissée",           f"{g['M_kphi']:.3f}",      "kN·m"],
                    ["θ au point M_max",        f"{g['theta_kphi']:.6f}",  "rad"],
                    ["M_max / θ",               f"{g['K_phi_brut']:.1f}",  "kN·m/rad"],
                    ["Facteur §A.5.2.3(6)",     "0,9",                     "—"],
                    ["Kφ,exp retenu",           f"{g['K_phi_exp']:.1f}",   "kN·m/rad"],
                ]
                tbl2 = doc.add_table(rows=len(kphi_rows)+1, cols=3)
                tbl2.style = "Table Grid"
                tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, h in enumerate(["Grandeur", "Valeur", "Unité"]):
                    cell = tbl2.rows[0].cells[i]
                    cell.paragraphs[0].clear()
                    run = cell.paragraphs[0].add_run(h)
                    run.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    set_cell_bg(cell, "2E75B6")
                for ri, row_d in enumerate(kphi_rows):
                    is_last = (ri == len(kphi_rows) - 1)
                    for ci, val in enumerate(row_d):
                        cell = tbl2.rows[ri+1].cells[ci]
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run(val)
                        run.font.name = "Arial"
                        run.font.size = Pt(10)
                        run.bold = is_last
                        set_cell_bg(cell, "FEF3C7" if is_last else ("EBF3FB" if ri % 2 == 1 else "FFFFFF"))

            # ── Sauvegarde ──
            buf_docx = BytesIO()
            doc.save(buf_docx)
            buf_docx.seek(0)

            st.download_button(
                "⬇ Télécharger le compte rendu (.docx)",
                buf_docx.getvalue(),
                "compte_rendu_Kphi.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            st.success("✅ Compte rendu généré avec succès !")
